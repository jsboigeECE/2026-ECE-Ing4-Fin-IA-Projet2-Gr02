from __future__ import annotations

import csv
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
RESULTS = ROOT / "results"
REPORTS = ROOT / "reports"
REPORTS.mkdir(parents=True, exist_ok=True)


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def round4(x: float) -> str:
    return f"{x:.4f}"


def round2(x: float) -> str:
    return f"{x:.2f}"


def add_heading_paragraph(doc: Document, text: str, size: int = 12, bold: bool = True) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def main() -> None:
    # Load key artifacts.
    if_metrics = load_json(RESULTS / "isolation_forest_metrics.json")
    if_best = load_json(RESULTS / "isolation_forest_best_threshold.json")
    ae_metrics = load_json(RESULTS / "autoencoder_detection_metrics.json")
    focal_summary = load_json(RESULTS / "focal_loss_summary.json")
    gnn_metrics = load_json(RESULTS / "gnn_metrics.json")
    gnn_best = load_json(RESULTS / "gnn_best_threshold.json")
    stream_summary = load_json(RESULTS / "streaming" / "stream_summary.json")

    doc = Document()

    doc.add_heading("Rapport Projet 2 - Detection de Fraude en Temps Reel", level=0)
    doc.add_paragraph(f"Date de generation: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("Dataset: Credit Card Fraud Detection (transactions cartes bancaires)")
    doc.add_paragraph("Fonction de cout utilisee: Cout = FP x 1 + FN x 25")

    doc.add_heading("1. Resume Executif", level=1)
    doc.add_paragraph(
        "Le projet a implemente et compare plusieurs approches de detection de fraude: "
        "Isolation Forest, Autoencoder, approches de reequilibrage (class weights, focal loss), "
        "Graph Neural Network (GNN) et simulation streaming temps reel."
    )
    doc.add_paragraph(
        "Les performances varient selon le compromis precision/recall et selon le contexte "
        "(batch offline vs streaming faible latence)."
    )

    doc.add_heading("2. Resultats par Methode", level=1)

    add_heading_paragraph(doc, "2.1 Isolation Forest")
    doc.add_paragraph(
        f"Metriques initiales: precision={round4(if_metrics['precision'])}, "
        f"recall={round4(if_metrics['recall'])}, AUPRC={round4(if_metrics['auprc'])}."
    )
    doc.add_paragraph(
        f"Apres optimisation de seuil (t={round4(if_best['best_threshold'])}): "
        f"precision={round4(if_best['precision'])}, recall={round4(if_best['recall'])}, "
        f"cout total={if_best['total_cost']:.0f}."
    )

    add_heading_paragraph(doc, "2.2 Autoencoder")
    doc.add_paragraph(
        f"Seuil de reconstruction: t={round4(ae_metrics['threshold'])}. "
        f"precision={round4(ae_metrics['precision'])}, "
        f"recall={round4(ae_metrics['recall'])}, AUPRC={round4(ae_metrics['auprc'])}."
    )
    doc.add_paragraph(
        f"Erreur moyenne normale={round4(ae_metrics['mean_error_normal'])}, "
        f"fraude={round4(ae_metrics['mean_error_fraud'])}."
    )

    add_heading_paragraph(doc, "2.3 Focal Loss (class imbalance)")
    table_focal = doc.add_table(rows=1, cols=5)
    hdr = table_focal.rows[0].cells
    hdr[0].text = "Modele"
    hdr[1].text = "Precision"
    hdr[2].text = "Recall"
    hdr[3].text = "AUPRC"
    hdr[4].text = "Cout"
    for row in focal_summary.get("results", []):
        cells = table_focal.add_row().cells
        cells[0].text = row["model"]
        cells[1].text = round4(row["precision"])
        cells[2].text = round4(row["recall"])
        cells[3].text = round4(row["auprc"])
        cells[4].text = f"{row['cost']:.0f}"

    add_heading_paragraph(doc, "2.4 Graph Neural Network (GNN)")
    doc.add_paragraph(
        f"Graphe: {gnn_metrics['num_nodes']} noeuds, {gnn_metrics['num_edges']} aretes."
    )
    doc.add_paragraph(
        f"Seuil cout-optimal: t={round4(gnn_best['cost_optimal']['threshold'])}, "
        f"precision={round4(gnn_best['cost_optimal']['precision'])}, "
        f"recall={round4(gnn_best['cost_optimal']['recall'])}, "
        f"cout={gnn_best['cost_optimal']['cost']:.0f}."
    )
    doc.add_paragraph(
        f"Seuil operationnel (precision >= {round2(gnn_best['operational_precision_constrained']['target_precision'])}): "
        f"t={round4(gnn_best['operational_precision_constrained']['threshold'])}, "
        f"precision={round4(gnn_best['operational_precision_constrained']['precision'])}, "
        f"recall={round4(gnn_best['operational_precision_constrained']['recall'])}, "
        f"cout={gnn_best['operational_precision_constrained']['cost']:.0f}."
    )

    comp = gnn_metrics.get("comparison_same_nodes", [])
    if comp:
        add_heading_paragraph(doc, "Comparaison sur les memes noeuds (GNN vs IF vs AE)", size=11)
        table_comp = doc.add_table(rows=1, cols=5)
        h = table_comp.rows[0].cells
        h[0].text = "Modele"
        h[1].text = "Precision"
        h[2].text = "Recall"
        h[3].text = "AUPRC"
        h[4].text = "Cout"
        for row in comp:
            c = table_comp.add_row().cells
            c[0].text = row["model"]
            c[1].text = round4(row["precision"])
            c[2].text = round4(row["recall"])
            c[3].text = round4(row["auprc"])
            c[4].text = f"{row['cost']:.0f}"

    doc.add_heading("3. Streaming Temps Reel (Phase 6)", level=1)
    doc.add_paragraph(
        f"Transactions traitees: {stream_summary['events_processed']} | "
        f"objectif latence: < {stream_summary['target_latency_ms']:.0f} ms"
    )
    doc.add_paragraph(
        f"Latence moyenne={round2(stream_summary['latency_avg_ms'])} ms, "
        f"p95={round2(stream_summary['latency_p95_ms'])} ms, "
        f"max={round2(stream_summary['latency_max_ms'])} ms."
    )
    doc.add_paragraph(
        f"Taux de transactions sous objectif latence: "
        f"{stream_summary['latency_under_target_rate'] * 100:.1f}%"
    )

    doc.add_heading("4. Conclusions", level=1)
    doc.add_paragraph(
        "Le projet est complet de bout en bout: pretraitement, modeles batch, optimisation de seuil, "
        "approche graphe et pipeline streaming faible latence."
    )
    doc.add_paragraph(
        "La priorite en production est de choisir un seuil operationnel selon la politique risque: "
        "maximiser la detection (recall) ou reduire les faux positifs (precision)."
    )

    out = REPORTS / "Rapport_Projet2_Detection_Fraude.docx"
    doc.save(out)
    print(str(out))


if __name__ == "__main__":
    main()
